# -*- coding: utf-8 -*-
"""
Exporter v3.4 — TXT + DOCX (SIMPLE-TNR) with OVERVIEW (Tree + Manifest)
- Tất cả văn bản trong DOCX: Times New Roman 14, không italic/shading, không code-style.
- Phần đầu tài liệu: PROJECT TREE (ASCII) + FILE MANIFEST (dir_rel/filename + size + status).
- Lọc phạm vi: --include-dirs, --include-names, --names-root-only.
- Chia nhỏ đầu ra (--split-chars), log CSV, report MD.
- Ví dụ: chỉ lấy Classes/**/*.cpp|.h + CMakeLists.txt (gốc).

pip install python-docx tqdm colorama
"""

import os, re, sys, csv, argparse
from pathlib import Path
from typing import List, Set, Tuple, Optional, Dict, Any
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

from tqdm import tqdm
try:
    import colorama
    colorama.just_fix_windows_console()
except Exception:
    pass

# --------- tên file đặc biệt (luôn cho phép) ----------
CODE_BASENAMES = {
    "CMakeLists.txt", "Makefile", "makefile", ".gitignore", ".gitattributes", "BUILD", "WORKSPACE"
}

# --------- preset gọn ----------
PRESETS = {
    "cpp": {
        "exts": {".c",".cc",".cpp",".cxx",".h",".hh",".hpp",".hxx",".m",".mm",".cmake"},
        "basenames": CODE_BASENAMES,
        "extra_exclude": set(),
    },
    "cocos": {
        "exts": {".c",".cc",".cpp",".cxx",".h",".hh",".hpp",".hxx",".m",".mm",".cmake",
                 ".glsl",".vert",".frag"},
        "basenames": CODE_BASENAMES,
        "extra_exclude": set(),
    },
}

EXCLUDE_DIRS_DEFAULT: Set[str] = {
    ".git",".svn",".hg",".idea",".vscode",".vs",".cache","__pycache__",".mypy_cache",".pytest_cache",
    "node_modules","dist","build","target","out","bin","obj","Library","Temp","Packages",".gradle",
    ".venv","venv",".nuget","Pods",".expo",".next",".angular","coverage"
}

SIZE_LIMIT_MB_DEFAULT = 5
DOCX_RUN_CHUNK = 100_000
PRINTABLE_WS = set(b"\t\n\r\f\v ")

# --------- utils ----------
def is_probably_text(sample: bytes) -> bool:
    if b"\x00" in sample:
        return False
    non_print = 0
    for b in sample:
        if 32 <= b <= 126:  # ASCII printable
            continue
        if b in PRINTABLE_WS:
            continue
        if b >= 128:        # UTF-8 multibyte
            continue
        non_print += 1
    return non_print / max(1, len(sample)) <= 0.30

def sanitize_for_docx(text: str) -> str:
    # bỏ control không hợp lệ với XML (trừ \t \n \r)
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)

def read_text_safely(p: Path, max_bytes: int) -> Tuple[str, str]:
    try:
        size = p.stat().st_size
    except Exception as e:
        return "error", f"[ERROR: cannot stat this file: {e}]\n"
    if size > max_bytes:
        return "skipped_large", f"[SKIPPED: file too large > {max_bytes//(1024*1024)}MB]\n"
    try:
        raw = p.read_bytes()
        if b"\x00" in raw:
            return "skipped_binary", "[SKIPPED: binary file (NUL found)]\n"
        if not is_probably_text(raw[: min(8192, len(raw))]):
            return "skipped_likely_binary", "[SKIPPED: likely binary]\n"
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("utf-8", errors="ignore")
        return "ok", text
    except Exception as e:
        return "error", f"[ERROR reading file: {e}]\n"

def should_exclude_dir(dirname: str, exclude_names: Set[str], include_hidden: bool) -> bool:
    name = os.path.basename(dirname)
    if not include_hidden and name.startswith('.'):
        return True
    return name in exclude_names

def path_matches_any_dir(path: Path, root: Path, include_dirs: List[str]) -> bool:
    rel = path.relative_to(root)
    parts = list(rel.parts)
    for d in include_dirs:
        d_parts = d.replace("\\","/").split("/")
        if parts[:len(d_parts)] == d_parts:
            return True
    return False

def should_include_file(
    path: Path,
    exts: Optional[Set[str]],
    include_hidden: bool,
    basenames: Set[str],
    root: Path,
    include_dirs: Optional[List[str]],
    include_names: Optional[List[str]],
    names_root_only: bool
) -> bool:
    base = path.name
    if not include_hidden and base.startswith('.'):
        return False

    # Ưu tiên phạm vi (nếu có)
    if include_dirs or include_names:
        ok = False
        if include_dirs and path_matches_any_dir(path, root, include_dirs):
            ok = True
        if include_names and base in set(include_names):
            if names_root_only:
                ok = ok or (path.parent.resolve() == root.resolve())
            else:
                ok = True
        if not ok:
            return False
        # Trong phạm vi rồi thì tiếp tục lọc theo đuôi hoặc tên đặc biệt
        if base in basenames:
            return True
        return (exts is None) or (path.suffix.lower() in exts)

    # Không chỉ định phạm vi: theo preset/--exts + basenames
    if base in basenames:
        return True
    if exts is None:
        return True
    return path.suffix.lower() in exts

# --------- DOCX (SIMPLE-TNR) ----------
def init_docx_simple() -> Document:
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(14)

    # Style nhẹ cho nhãn "Folder/File" (vẫn TNR 14, chỉ đậm)
    if 'LabelBold' not in [s.name for s in doc.styles]:
        st = doc.styles.add_style('LabelBold', WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = doc.styles['LabelBold']
    st.font.name = 'Times New Roman'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    st.font.size = Pt(14)
    st.font.bold = True
    return doc

def add_plain_paragraph(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(14)
    r.bold = bool(bold)

def add_text_block(doc: Document, text: str):
    """Chèn nội dung văn bản (code) như đoạn thường, TNR 14, giữ \\n."""
    safe = sanitize_for_docx(text)
    n = len(safe)
    start = 0
    p = doc.add_paragraph()
    while start < n:
        part = safe[start : start + DOCX_RUN_CHUNK]
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(14)
        start += DOCX_RUN_CHUNK

# --------- OVERVIEW (TREE + MANIFEST) ----------
def collect_dir_hierarchy(grouped: Dict[str, List[Dict[str, Any]]]) -> Tuple[Dict[str, List[str]], Dict[str, List[str]]]:
    """
    Trả về:
      - children_dirs[parent_dir] = [child_dir, ...]
      - files_in_dir[dir] = [filename, ...]
    Bổ sung cả các thư mục tổ tiên để tree không bị thiếu mắt xích.
    """
    # Tập hợp tất cả thư mục đã có + tổ tiên
    all_dirs = set(grouped.keys())
    for d in list(all_dirs):
        if d == ".":
            continue
        cur = d
        while True:
            par = os.path.dirname(cur) or "."
            if par not in all_dirs:
                all_dirs.add(par)
            if par == ".":
                break
            cur = par

    children_dirs: Dict[str, List[str]] = defaultdict(list)
    files_in_dir: Dict[str, List[str]] = defaultdict(list)

    for d in all_dirs:
        files_in_dir[d] = []

    for d, items in grouped.items():
        for it in items:
            files_in_dir[d].append(it["fname"])

    # Sắp xếp tên thư mục con
    for d in sorted(all_dirs):
        if d == ".":
            par = None
        else:
            par = os.path.dirname(d) or "."
        if par is not None and d != ".":
            children_dirs[par].append(d)

    for k in children_dirs:
        # chỉ giữ tên con là basename, nhưng lưu kèm đường dẫn đầy đủ để render đúng thứ tự
        children_dirs[k] = sorted(children_dirs[k], key=lambda x: x)

    # Sort files
    for d in files_in_dir:
        files_in_dir[d] = sorted(files_in_dir[d], key=lambda x: x.lower())

    return children_dirs, files_in_dir

def render_ascii_tree(children_dirs: Dict[str, List[str]], files_in_dir: Dict[str, List[str]]) -> str:
    """
    Vẽ ASCII tree kiểu:
    ./
    ├── CMakeLists.txt
    └── Classes
        ├── core
        │   └── AppDelegate.cpp
        ...
    """
    lines: List[str] = []

    def name_of(path_rel: str) -> str:
        return "." if path_rel == "." else os.path.basename(path_rel)

    def recurse(dir_rel: str, prefix: str):
        # Files trước
        files = files_in_dir.get(dir_rel, [])
        dirs  = children_dirs.get(dir_rel, [])
        total = len(files) + len(dirs)
        idx = 0

        # Files
        for i, fn in enumerate(files):
            idx += 1
            is_last = (idx == total)
            lines.append(f"{prefix}{'└──' if is_last and not dirs else '├──'} {fn}")

        # Dirs
        for j, child in enumerate(dirs):
            is_last_dir = (j == len(dirs) - 1)
            connector = "└──" if is_last_dir else "├──"
            lines.append(f"{prefix}{connector} {name_of(child)}")
            new_prefix = prefix + ("    " if is_last_dir else "│   ")
            recurse(child, new_prefix)

    # Root
    lines.append("./")
    recurse(".", "")
    return "\n".join(lines) + "\n"

def render_manifest(grouped: Dict[str, List[Dict[str, Any]]]) -> str:
    """
    Liệt kê đường dẫn + size + status (ok/skipped/error).
    """
    rows: List[str] = []
    for d in sorted(grouped.keys(), key=lambda s: (0 if s=="." else 1, s)):
        for it in sorted(grouped[d], key=lambda x: x["fname"].lower()):
            size = it.get("size", -1)
            st   = it.get("status","")
            path = f"{d}/{it['fname']}" if d != "." else it['fname']
            rows.append(f"- {path} — {size} bytes — {st}")
    return "\n".join(rows) + ("\n" if rows else "")

# --------- worker ----------
def process_one_file(args):
    root, fullpath, max_bytes = args
    rel_dir = os.path.relpath(fullpath.parent, root)
    rel_dir = "." if rel_dir == "." else rel_dir.replace("\\","/")
    fname = fullpath.name
    try:
        size = fullpath.stat().st_size
    except Exception:
        size = -1
    status, payload = read_text_safely(fullpath, max_bytes)
    return {"dir_rel": rel_dir, "fname": fname, "size": size, "status": status, "text": payload}

# --------- ghi ra (có chia nhỏ) ----------
def write_outputs_split(
    grouped: Dict[str, List[Dict[str, Any]]],
    out_txt_base: Path,
    out_docx_base: Path,
    split_chars: int,
    add_overview: bool
):
    # TXT
    txt_parts: List[Path] = []
    txt_idx = 1
    txt_chars = 0
    if split_chars and split_chars > 0:
        txt_path = out_txt_base.with_name(f"{out_txt_base.stem}_part{txt_idx:02d}{out_txt_base.suffix}")
    else:
        txt_path = out_txt_base
    ftxt = txt_path.open("w", encoding="utf-8", newline="\n")
    txt_parts.append(txt_path)

    # DOCX
    docx_parts: List[Path] = []
    docx_idx = 1
    docx_chars = 0
    doc = init_docx_simple()
    if split_chars:
        first_docx = out_docx_base.with_name(f"{out_docx_base.stem}_part{docx_idx:02d}{out_docx_base.suffix}")
    else:
        first_docx = out_docx_base
    docx_parts.append(first_docx)

    def txt_write(s: str):
        nonlocal txt_chars, ftxt, txt_idx, txt_path
        ftxt.write(s)
        txt_chars += len(s)
        if split_chars and txt_chars >= split_chars:
            ftxt.close()
            txt_idx += 1
            txt_path = out_txt_base.with_name(f"{out_txt_base.stem}_part{txt_idx:02d}{out_txt_base.suffix}")
            txt_parts.append(txt_path)
            ftxt = txt_path.open("w", encoding="utf-8", newline="\n")
            txt_chars = 0

    def docx_rotate_if_needed():
        nonlocal doc, docx_chars, docx_idx
        if split_chars and docx_chars >= split_chars:
            doc.save(str(docx_parts[-1]))
            docx_idx += 1
            new_path = out_docx_base.with_name(f"{out_docx_base.stem}_part{docx_idx:02d}{out_docx_base.suffix}")
            docx_parts.append(new_path)
            doc = init_docx_simple()
            docx_chars = 0

    # ---------- OVERVIEW (only at the very beginning / first part) ----------
    if add_overview:
        children_dirs, files_in_dir = collect_dir_hierarchy(grouped)
        tree_text = render_ascii_tree(children_dirs, files_in_dir)
        manifest_text = render_manifest(grouped)

        # TXT overview
        txt_write("===== PROJECT TREE =====\n")
        txt_write(tree_text + "\n")
        txt_write("===== FILE MANIFEST =====\n")
        txt_write(manifest_text + "\n\n")

        # DOCX overview
        add_plain_paragraph(doc, "===== PROJECT TREE =====", bold=True)
        docx_chars += len("===== PROJECT TREE =====")
        add_text_block(doc, tree_text)
        docx_chars += len(tree_text)
        docx_rotate_if_needed()

        add_plain_paragraph(doc, "===== FILE MANIFEST =====", bold=True)
        docx_chars += len("===== FILE MANIFEST =====")
        add_text_block(doc, manifest_text)
        docx_chars += len(manifest_text)
        docx_rotate_if_needed()

        # Ngăn cách
        add_plain_paragraph(doc, "", bold=False)
        txt_write("\n")

    # ---------- Nội dung theo Folder/File ----------
    for dir_rel in sorted(grouped.keys(), key=lambda s: (0 if s=="." else 1, s)):
        # TXT
        txt_write(f"Folder {dir_rel}\n")
        # DOCX (nhãn gọn)
        add_plain_paragraph(doc, f"Folder {dir_rel}", bold=True)
        docx_chars += len(f"Folder {dir_rel}")
        docx_rotate_if_needed()

        for item in sorted(grouped[dir_rel], key=lambda x: x["fname"].lower()):
            fname = item["fname"]
            content = item["text"]

            txt_write(f"File {fname}\n")
            txt_write(content.rstrip("\n") + "\n\n")

            add_plain_paragraph(doc, f"File {fname}", bold=True)
            docx_chars += len(f"File {fname}")
            add_text_block(doc, content)
            docx_chars += len(content)
            docx_rotate_if_needed()

        txt_write("\n")

    ftxt.close()
    doc.save(str(docx_parts[-1]))
    return txt_parts, docx_parts

# --------- export chính ----------
def export_project(
    root: Path,
    out_txt: Path,
    out_docx: Path,
    log_csv: Path,
    report_md: Path,
    exts: Optional[Set[str]],
    basenames: Set[str],
    exclude_dirs: Set[str],
    include_hidden: bool,
    size_limit_mb: int,
    workers: int,
    split_chars: int,
    include_dirs: Optional[List[str]],
    include_names: Optional[List[str]],
    names_root_only: bool,
    overview: bool
):
    max_bytes = size_limit_mb * 1024 * 1024
    root = root.resolve()

    # Quét ứng viên
    candidates: List[Path] = []
    for dirpath, dirnames, filenames in os.walk(root, topdown=True):
        dirnames[:] = [d for d in sorted(dirnames)
                       if not should_exclude_dir(os.path.join(dirpath,d), exclude_dirs, include_hidden)]
        for fname in sorted(filenames):
            p = Path(dirpath)/fname
            if should_include_file(
                p, exts, include_hidden, basenames, root, include_dirs, include_names, names_root_only
            ):
                candidates.append(p)

    total = len(candidates)
    if total == 0:
        print("Không có file phù hợp. Kiểm tra --include-dirs/--include-names/--exts.")
        return

    grouped: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    stats = {"ok":0,"skip":0,"error":0}
    log_rows: List[List[Any]] = []

    with ThreadPoolExecutor(max_workers=max(1, workers)) as ex:
        futures = {ex.submit(process_one_file, (root, p, max_bytes)): p for p in candidates}
        with tqdm(total=total, desc="Đang đọc & xử lý", unit="file") as bar:
            for fut in as_completed(futures):
                p = futures[fut]
                try:
                    item = fut.result()
                except Exception as e:
                    rel_dir = os.path.relpath(p.parent, root)
                    rel_dir = "." if rel_dir=="." else rel_dir.replace("\\","/")
                    log_rows.append([rel_dir, p.name, -1, "error", f"worker exception: {e}"])
                    stats["error"] += 1
                    bar.update(1)
                    bar.set_postfix(ok=stats["ok"], skip=stats["skip"], err=stats["error"])
                    tqdm.write(f"[ERROR] {p}")
                    continue

                grouped[item["dir_rel"]].append(item)
                st = item["status"]; size = item["size"]; note = "" if st=="ok" else item["text"].strip()
                if st=="ok": stats["ok"] += 1
                elif st.startswith("skipped"): stats["skip"] += 1
                else: stats["error"] += 1
                log_rows.append([item["dir_rel"], item["fname"], size, st, note])

                bar.update(1)
                bar.set_postfix(ok=stats["ok"], skip=stats["skip"], err=stats["error"])
                bar.set_description(f"Đang xử lý: {item['fname'][:40]}")

    # Ghi TXT/DOCX
    with tqdm(total=1, desc="Đang ghi TXT/DOCX") as bar2:
        txt_parts, docx_parts = write_outputs_split(grouped, out_txt, out_docx, split_chars, add_overview=overview)
        bar2.update(1)

    # Log CSV
    with log_csv.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f); w.writerow(["dir_rel","filename","size_bytes","status","note"])
        w.writerows(log_rows)

    # Report
    rpt = []
    rpt.append("# Export Report\n")
    rpt.append(f"- Root: `{root}`")
    rpt.append(f"- Tổng file xét: **{total}**")
    rpt.append(f"- OK: **{stats['ok']}** | SKIP: **{stats['skip']}** | ERROR: **{stats['error']}**\n")
    rpt.append("## Output files")
    rpt.append("- TXT: " + ", ".join(f"{p.name}" for p in txt_parts))
    rpt.append("- DOCX: " + ", ".join(f"{p.name}" for p in docx_parts))
    if stats["error"] or stats["skip"]:
        rpt.append("\n## Mục skip/lỗi (tối đa 50)")
        c = 0
        for row in log_rows:
            if row[3] != "ok":
                rpt.append(f"- `{row[0]}/{row[1]}` — **{row[3]}** — {row[4]}")
                c += 1
                if c >= 50: break
    report_md.write_text("\n".join(rpt), encoding="utf-8")

    print("\n✅ Hoàn tất.")
    print("TXT :", ", ".join(str(p) for p in txt_parts))
    print("DOCX:", ", ".join(str(p) for p in docx_parts))
    print("LOG :", log_csv)
    print("REPORT:", report_md)
    print(f"Tổng: {total} | ok={stats['ok']} skip={stats['skip']} err={stats['error']}")

# --------- CLI ----------
def parse_args(argv: List[str]) -> argparse.Namespace:
    ap = argparse.ArgumentParser(description="Xuất dự án ra TXT & DOCX (SIMPLE-TNR).")
    ap.add_argument("--root", default=".", help="Thư mục gốc (mặc định: .)")
    ap.add_argument("--out-txt", default="project_dump.txt", help="TXT đầu ra (có thể _partNN)")
    ap.add_argument("--out-docx", default="project_dump.docx", help="DOCX đầu ra (có thể _partNN)")
    ap.add_argument("--log-csv", default="export_log.csv", help="Log CSV")
    ap.add_argument("--report-md", default="export_report.md", help="Báo cáo Markdown")
    ap.add_argument("--size-limit-mb", type=int, default=SIZE_LIMIT_MB_DEFAULT, help="Giới hạn kích cỡ file đọc (MB)")
    ap.add_argument("--include-hidden", action="store_true", help="Bao gồm file/thư mục ẩn")
    ap.add_argument("--all", action="store_true", help="Lấy tất cả file văn bản (trừ nhị phân)")
    ap.add_argument("--exts", nargs="*", default=None, help="Chỉ lấy các đuôi chỉ định (bỏ qua nếu --all)")
    ap.add_argument("--preset", choices=list(PRESETS.keys()), default="cpp", help="Chọn preset")
    ap.add_argument("--add-ext", nargs="*", default=None, help="Bổ sung đuôi (vd: --add-ext .proto .conf)")
    ap.add_argument("--no-default-exclude", action="store_true", help="Không loại trừ thư mục rác mặc định")
    ap.add_argument("--workers", type=int, default=max(1, (os.cpu_count() or 4)//2), help="Số luồng đọc")
    ap.add_argument("--split-chars", type=int, default=0, help="Giới hạn kí tự mỗi phần; >0 sẽ tách _partNN")
    # Phạm vi
    ap.add_argument("--include-dirs", nargs="*", default=None, help="Chỉ lấy các thư mục (prefix, vd: Classes)")
    ap.add_argument("--include-names", nargs="*", default=None, help="Thêm các file theo tên (vd: CMakeLists.txt)")
    ap.add_argument("--names-root-only", action="store_true", help="Các --include-names chỉ áp dụng ở root")
    # Overview
    ap.add_argument("--no-overview", action="store_true", help="Tắt phần PROJECT TREE + FILE MANIFEST ở đầu")
    return ap.parse_args(argv)

def main():
    args = parse_args(sys.argv[1:])
    root = Path(args.root).resolve()
    if not root.exists() or not root.is_dir():
        print(f"[ERROR] Root '{root}' không tồn tại hoặc không phải thư mục.")
        sys.exit(2)

    # exts/basenames theo ưu tiên: --all > --exts > --preset
    basenames: Set[str] = set()
    extra_exclude: Set[str] = set()
    if args.all:
        exts = None
        basenames = CODE_BASENAMES
    else:
        if args.exts:
            exts = set(args.exts)
            basenames = CODE_BASENAMES
        else:
            preset = PRESETS[args.preset]
            exts = set(preset["exts"])
            basenames = set(preset["basenames"])
            extra_exclude = set(preset.get("extra_exclude", set()))

    if args.add_ext:
        exts = (set() if exts is None else set(exts)) | set(args.add_ext)

    exclude = set() if args.no_default_exclude else (set(EXCLUDE_DIRS_DEFAULT) | extra_exclude)

    out_txt = root / args.out_txt
    out_docx = root / args.out_docx
    log_csv = root / args.log_csv
    report_md = root / args.report_md

    print("=== CONFIG SUMMARY ===")
    print(f"Root           : {root}")
    print(f"Preset         : {args.preset}")
    print(f"All            : {args.all}")
    print(f"Exts count     : {'ALL (text-like)' if exts is None else len(exts)}")
    print(f"Include dirs   : {args.include_dirs}")
    print(f"Include names  : {args.include_names} (root-only={args.names_root_only})")
    print(f"Split chars    : {args.split_chars}")
    print(f"Overview       : {not args.no_overview}")
    print("======================")

    export_project(
        root=root,
        out_txt=out_txt,
        out_docx=out_docx,
        log_csv=log_csv,
        report_md=report_md,
        exts=exts,
        basenames=basenames,
        exclude_dirs=exclude,
        include_hidden=args.include_hidden,
        size_limit_mb=max(1, args.size_limit_mb),
        workers=max(1, args.workers),
        split_chars=max(0, args.split_chars),
        include_dirs=args.include_dirs,
        include_names=args.include_names,
        names_root_only=args.names_root_only,
        overview=(not args.no_overview),
    )

if __name__ == "__main__":
    main()
